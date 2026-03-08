import os
from docx import Document
from typing import Dict, Any

class WordGeneratorService:
    def __init__(self, templates_dir: str = "./templates"):
        self.templates_dir = templates_dir
        if not os.path.exists(self.templates_dir):
            os.makedirs(self.templates_dir)

    def generate_document(self, template_name: str, meeting_data: Dict[str, Any], output_path: str) -> str:
        """
        Genera un acta en Word incrustando los datos extraídos por Groq.
        En un entorno real, usaría una plantilla existente y reemplazaría variables (ej {{Resumen}}).
        Por simplicidad del Agente, crearemos un documento desde cero estructurándolo.
        """
        # Si hubiera plantilla usaríamos: doc = Document(f"{self.templates_dir}/{template_name}.docx")
        doc = Document()
        
        # Título
        doc.add_heading(f"Acta de Reunión: {meeting_data.get('title', 'Sin Título')}", 0)
        
        doc.add_heading("Resumen Ejecutivo", level=1)
        doc.add_paragraph(meeting_data.get("summary", "Sin resumen"))
        
        doc.add_heading("Decisiones Clave", level=1)
        doc.add_paragraph(meeting_data.get("decisions", "Ninguna decisión registrada"))
        
        doc.add_heading("Riesgos Identificados", level=1)
        doc.add_paragraph(meeting_data.get("risks", "Ningún riesgo detectado"))
        
        doc.add_heading("Acuerdos", level=1)
        doc.add_paragraph(meeting_data.get("agreements", "Ningún acuerdo"))
        
        doc.add_heading("Tareas y Compromisos (Action Items)", level=1)
        
        items = meeting_data.get("action_items", [])
        if items:
            for item in items:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{item['title']} - {item['owner_name']} ").bold = True
                p.add_run(f"(Vence: {item.get('due_date', 'N/A')})\n")
                p.add_run(item.get("description", ""))
        else:
            doc.add_paragraph("No se detectaron tareas.")
            
        # Guardar Documento
        doc.save(output_path)
        return output_path
