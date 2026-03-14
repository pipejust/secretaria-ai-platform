import json
import asyncio
from sqlmodel import Session, select
from database import engine
from models import MeetingSession

# Create a mock session entry if needed. Or just test with Session 3 if it exists.
def run():
    with Session(engine) as db:
        s = db.get(MeetingSession, 3)
        if not s:
            print("Session 3 not found.")
            return

        import datetime
        print(s.title)
        
        # Test DocXTPL mapping
        from docxtpl import DocxTemplate
        import io
        
        # Generate dummy docxtpl
        tpl = DocxTemplate()
        from docx import Document
        doc = Document()
        doc.add_paragraph("Test: {{ title }}")
        doc.add_paragraph("{%p for a in attendees %}")
        doc.add_paragraph("   - {{ a.name }} / {{ a.role }}")
        doc.add_paragraph("{%p endfor %}")
        doc.save("test_dummy.docx")
        
        tpl = DocxTemplate("test_dummy.docx")
        
        attendees = [{"name": "Felipe Cortes", "role": "Líder"}]
        context = {
            "title": s.title,
            "attendees": attendees
        }
        tpl.render(context)
        tpl.save("test_output.docx")
        print("Done!")

run()
