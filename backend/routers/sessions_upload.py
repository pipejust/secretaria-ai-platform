from fastapi import APIRouter, UploadFile, File, Form, Depends, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from typing import Optional
from database import get_session
from sqlmodel import Session
from models import MeetingSession
import datetime
import uuid
import io
import docx

router = APIRouter(
    prefix="/api/sessions",
    tags=["Sessions"]
)

@router.get("/")
def get_sessions(db: Session = Depends(get_session)):
    """Fetch all meeting sessions (Actas) from the database."""
    from sqlmodel import select
    sessions = db.exec(select(MeetingSession).order_by(MeetingSession.id.desc())).all()
    return sessions

@router.get("/{session_id}")
def get_session_details(session_id: int, db: Session = Depends(get_session)):
    """Fetch a specific meeting session and its related action items."""
    from sqlmodel import select
    from models import ActionItem
    
    session_obj = db.get(MeetingSession, session_id)
    if not session_obj:
        raise HTTPException(status_code=404, detail="Session not found")
        
    action_items = db.exec(select(ActionItem).where(ActionItem.session_id == session_id)).all()
    # Podríamos crear un Pydantic model response, pero dict/jsonable encoder lo maneja bien
    return {
        "session": session_obj,
        "action_items": action_items
    }

class SessionUpdate(BaseModel):
    raw_summary: Optional[str] = None
    raw_transcript: Optional[str] = None
    processed_decisions: Optional[str] = None
    processed_risks: Optional[str] = None
    processed_agreements: Optional[str] = None

@router.post("/{session_id}/regenerate_tasks")
async def regenerate_tasks_from_transcript(session_id: int, db: Session = Depends(get_session)):
    from models import ActionItem, ProjectContact
    from services.groq_service import GroqService
    from sqlmodel import delete

    session_obj = db.get(MeetingSession, session_id)
    if not session_obj:
        raise HTTPException(status_code=404, detail="Session not found")
    
    if not session_obj.raw_transcript:
        raise HTTPException(status_code=400, detail="No transcript available to regenerate tasks from.")

    # 1. Obtenemos Contactos
    project_contacts = []
    if session_obj.project_id:
        from sqlmodel import select
        db_contacts = db.exec(select(ProjectContact).where(ProjectContact.project_id == session_obj.project_id)).all()
        project_contacts = [{"name": c.name, "email": c.email, "role": c.role} for c in db_contacts]

    # 2. Llamamos a Groq
    groq_svc = GroqService()
    structured_data = await groq_svc.process_transcript(session_obj.raw_transcript, project_contacts)

    # 3. Borramos Tareas Anteriores
    db.exec(delete(ActionItem).where(ActionItem.session_id == session_id))
    db.commit()

    # 4. Insertamos Nuevas
    action_items_data = structured_data.get("action_items", [])
    new_items_output = []
    for item_data in action_items_data:
        action_item = ActionItem(
            session_id=session_id,
            owner_name=item_data.get("owner_name", "Unknown"),
            owner_email=item_data.get("owner_email", ""),
            title=item_data.get("title", "Tarea sin título"),
            description=item_data.get("description", ""),
            due_date=item_data.get("due_date"),
            is_approved=False
        )
        db.add(action_item)
        db.commit()
        db.refresh(action_item)
        # Convertimos para el output JSON dict
        new_items_output.append({
            "id": action_item.id,
            "session_id": action_item.session_id,
            "owner_name": action_item.owner_name,
            "owner_email": action_item.owner_email,
            "title": action_item.title,
            "description": action_item.description,
            "due_date": str(action_item.due_date) if action_item.due_date else None,
            "is_approved": action_item.is_approved,
            "selected": False
        })

    return {"status": "success", "action_items": new_items_output}

@router.put("/{session_id}")
def update_session_content(session_id: int, payload: SessionUpdate, db: Session = Depends(get_session)):
    """Manually update the text content of a curated session."""
    session_obj = db.get(MeetingSession, session_id)
    if not session_obj:
        raise HTTPException(status_code=404, detail="Session not found")
        
    if payload.raw_summary is not None:
        session_obj.raw_summary = payload.raw_summary
    if payload.raw_transcript is not None:
        session_obj.raw_transcript = payload.raw_transcript
    if payload.processed_decisions is not None:
        session_obj.processed_decisions = payload.processed_decisions
    if payload.processed_risks is not None:
        session_obj.processed_risks = payload.processed_risks
    if payload.processed_agreements is not None:
        session_obj.processed_agreements = payload.processed_agreements
        
    db.add(session_obj)
    db.commit()
    db.refresh(session_obj)
    return {"status": "success", "message": "Manual edits saved successfully"}

@router.put("/action_items/{item_id}")
def update_action_item_email(item_id: int, owner_email: str = Form(...), db: Session = Depends(get_session)):
    """Update the owner email of an action item manually."""
    from models import ActionItem
    item = db.get(ActionItem, item_id)
    if not item:
        raise HTTPException(status_code=404, detail="Action Item not found")
    
    item.owner_email = owner_email
    db.add(item)
    db.commit()
    db.refresh(item)
    return {"status": "success", "message": "Email actualizado", "item": item}


@router.post("/upload")
async def upload_manual_session(
    title: str = Form(...),
    file: UploadFile = File(...),
    session: Session = Depends(get_session)
):
    try:
        # Extraer el contenido o guardar el archivo
        content = await file.read()
        file_size = len(content)
        
        # Crear una nueva sesión
        import uuid
        new_session = MeetingSession(
            fireflies_id=f"manual_{uuid.uuid4()}",
            title=title,
            date=datetime.datetime.utcnow().isoformat() + "Z",
            video_url=f"manual_upload_{file.filename}",
            raw_transcript=f"Uploaded {file.filename} manually.",  # Changed to raw_transcript to match model
            status="pending",
            raw_summary="",
            processed_decisions="",
            processed_risks="",
            processed_agreements=""
        )
        
        session.add(new_session)
        session.commit()
        session.refresh(new_session)
        
        return {"status": "success", "session_id": new_session.id, "message": "Archivo subido exitosamente."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

class DispatchEmailsRequest(BaseModel):
    action_item_ids: list[int]

@router.post("/{session_id}/dispatch_emails")
async def dispatch_emails(session_id: int, request: DispatchEmailsRequest, db: Session = Depends(get_session)):
    """Dispatch emails for the selected action items."""
    from models import ActionItem
    from services.email_service import EmailService
    
    session_obj = db.get(MeetingSession, session_id)
    if not session_obj:
        raise HTTPException(status_code=404, detail="Session not found")
        
    email_service = EmailService()
    results = []
    
    for item_id in request.action_item_ids:
        item = db.get(ActionItem, item_id)
        if not item or item.session_id != session_id:
            continue
            
        if not item.owner_email:
            results.append({"id": item_id, "status": "failed", "reason": "No email provided"})
            continue
            
        try:
            # Here we will call the new email service method sending the HTML template
            await email_service.send_action_item_email(
                to_email=item.owner_email,
                owner_name=item.owner_name,
                task_title=item.title,
                task_description=item.description,
                project_name=session_obj.title # or the actual project name
            )
            results.append({"id": item_id, "status": "success"})
        except Exception as e:
            results.append({"id": item_id, "status": "failed", "reason": str(e)})
            
    return {"status": "success", "results": results}

class DispatchPlatformsRequest(BaseModel):
    action_item_ids: list[int]

@router.post("/{session_id}/dispatch_platforms")
async def dispatch_platforms(session_id: int, request: DispatchPlatformsRequest, db: Session = Depends(get_session)):
    """Dispatch selected action items to configured platforms (ClickUp, Trello, Jira)."""
    from models import ActionItem, Routing
    import json
    from sqlmodel import select
    from services.integrations.trello import TrelloIntegrationService
    from services.integrations.jira import JiraIntegrationService
    from services.integrations.clickup import ClickUpIntegrationService

    session_obj = db.get(MeetingSession, session_id)
    if not session_obj:
        raise HTTPException(status_code=404, detail="Session not found")

    if not session_obj.project_id:
        raise HTTPException(status_code=400, detail="Cannot dispatch: Meeting is not related to any project routing.")
        
    routings = db.exec(select(Routing).where(Routing.project_id == session_obj.project_id)).all()
    if not routings:
        raise HTTPException(status_code=400, detail="Project has no configured routings.")

    results = []
    for item_id in request.action_item_ids:
        item = db.get(ActionItem, item_id)
        if not item or item.session_id != session_id:
            continue
        
        item_success = False
        for routing in routings:
            config = json.loads(routing.destination_config or '{}')
            dest_type = routing.destination_type.lower()
            
            try:
                if "trello" in dest_type:
                    trello_service = TrelloIntegrationService("mock_key", "mock_token") # TODO use Settings
                    await trello_service.create_card(config.get("board_id"), config.get("list_id"), item.title, item.description, item.due_date)
                    item_success = True
                elif "jira" in dest_type:
                    jira_service = JiraIntegrationService("mock_domain", "mock@email.com", "mock_token") 
                    await jira_service.create_issue(config.get("project_key"), item.title, item.description)
                    item_success = True
                elif "clickup" in dest_type:
                    clickup_service = ClickUpIntegrationService("mock_token") 
                    await clickup_service.create_task(config.get("list_id"), item.title, item.description)
                    item_success = True
            except Exception as e:
                pass # Proceed to next routing iteration
                
        if item_success:
            results.append({"id": item_id, "status": "success"})
        else:
            results.append({"id": item_id, "status": "failed"})

    return {"status": "success", "results": results}

@router.get("/{session_id}/export/word")
def export_word(session_id: int, db: Session = Depends(get_session)):
    """Generate and return a Microsoft Word (.docx) document with the meeting details."""
    from models import ActionItem
    from sqlmodel import select
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    import io

    session_obj = db.get(MeetingSession, session_id)
    if not session_obj:
        raise HTTPException(status_code=404, detail="Session not found")

    action_items = db.exec(select(ActionItem).where(ActionItem.session_id == session_id)).all()

    doc = Document()

    # Title
    title_run = doc.add_heading(level=0).add_run("Acta de Reunión")
    title_run.font.color.rgb = RGBColor(79, 70, 229) # Branding Indigo
    
    # Meta
    doc.add_paragraph(f"Proyecto / Sesión: {session_obj.title}", style='Intense Quote')
    doc.add_paragraph(f"Fecha: {session_obj.date}")
    doc.add_paragraph(f"Estado: {session_obj.status.capitalize()}")
    doc.add_paragraph()

    # Sections
    if session_obj.raw_summary:
        doc.add_heading('Resumen Ejecutivo', level=1)
        doc.add_paragraph(session_obj.raw_summary)

    if session_obj.processed_decisions:
        doc.add_heading('Decisiones Clave', level=1)
        doc.add_paragraph(session_obj.processed_decisions)

    if session_obj.processed_risks:
        doc.add_heading('Riesgos Identificados', level=1)
        doc.add_paragraph(session_obj.processed_risks)

    if session_obj.processed_agreements:
        doc.add_heading('Acuerdos', level=1)
        doc.add_paragraph(session_obj.processed_agreements)

    # Action Items Table
    doc.add_heading('Tareas (Action Items)', level=1)
    
    if action_items:
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Responsable'
        hdr_cells[1].text = 'Tarea'
        hdr_cells[2].text = 'Descripción'
        hdr_cells[3].text = 'Vencimiento'
        
        for item in action_items:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{item.owner_name} ({item.owner_email})"
            row_cells[1].text = item.title
            row_cells[2].text = item.description or ""
            row_cells[3].text = item.due_date or "Sin fecha"
    else:
        doc.add_paragraph("No se detectaron tareas para esta sesión.")
        
    # Salvar el documento a un buffer en memoria
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    headers = {
        'Content-Disposition': f'attachment; filename="Acta_{session_obj.id}_{session_obj.title[:20]}.docx"'
    }

    return Response(content=buffer.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)
